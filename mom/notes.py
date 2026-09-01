"""Phase 1 — the meeting note becomes structured JSON.

Two readers, chosen by file extension, produce the same shape: a list of pages of
text plus any tables found on them. Only the reader knows about a file format; the
parse step below sees pages either way.

The standalone version read PDFs only, while its "Self Notes" upload accepted .docx —
a .docx upload reached PyMuPDF and crashed. The dispatch table is the fix.
"""
from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, List

from logger import get_logger
from mom.llm import JsonCaller

log = get_logger(__name__)


@dataclass(frozen=True)
class NotePage:
    """One page of the note: its running text, plus tables as markdown."""

    page: int
    text: str
    tables: tuple[str, ...] = ()

    def as_prompt_block(self) -> str:
        return self.text + ("\n\n" + "\n\n".join(self.tables) if self.tables else "")


class UnsupportedNoteFormat(ValueError):
    """The uploaded note is not a format the pipeline can read."""


# ── readers ───────────────────────────────────────────────────────────────────


def _rows_to_markdown(rows: List[List[str]]) -> str | None:
    """A table as a markdown pipe block, or None when there is nothing in it."""
    cleaned = [[str(cell or "").strip() for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(row)]
    if not cleaned:
        return None
    header, *body = cleaned
    lines = ["| " + " | ".join(header) + " |", "|" + "|".join(["---"] * len(header)) + "|"]
    lines += ["| " + " | ".join(row) + " |" for row in body]
    return "\n".join(lines)


def read_pdf(path: Path) -> List[NotePage]:
    """Text and tables per page, via PyMuPDF."""
    import pymupdf

    pages: List[NotePage] = []
    document = pymupdf.open(str(path))
    try:
        for number, page in enumerate(document, start=1):
            tables = []
            try:
                for table in page.find_tables():
                    markdown = _rows_to_markdown(table.extract() or [])
                    if markdown:
                        tables.append(markdown)
            except AttributeError:
                pass  # find_tables() needs pymupdf >= 1.23
            pages.append(
                NotePage(page=number, text=page.get_text("text") or "", tables=tuple(tables))
            )
    finally:
        document.close()
    return pages


def read_docx(path: Path) -> List[NotePage]:
    """A Word document has no pages until it is rendered, so it reads as one page."""
    from docx import Document

    document = Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    tables = []
    for table in document.tables:
        markdown = _rows_to_markdown([[cell.text for cell in row.cells] for row in table.rows])
        if markdown:
            tables.append(markdown)
    return [NotePage(page=1, text=text, tables=tuple(tables))]


READERS: Dict[str, Callable[[Path], List[NotePage]]] = {
    ".pdf": read_pdf,
    ".docx": read_docx,
}


def read_note(path: Path) -> List[NotePage]:
    """The note's pages, whichever supported format it arrived in."""
    reader = READERS.get(Path(path).suffix.lower())
    if reader is None:
        raise UnsupportedNoteFormat(
            f"Cannot read {Path(path).name}: supported formats are "
            f"{', '.join(sorted(READERS))}."
        )
    return reader(Path(path))


# ── parse ─────────────────────────────────────────────────────────────────────


def build_parse_prompt(pages: List[NotePage]) -> str:
    """The extraction prompt. Client, date and attendees come from the deck instead."""
    body = "\n\n--- PAGE BREAK ---\n\n".join(page.as_prompt_block() for page in pages)
    return (
        "You are extracting structured information from a QBR meeting note for Marsh ICG.\n"
        "The FORMAT IS VARIABLE — formal Marsh template, informal bullets, Zoom AI summary,\n"
        "party-labelled sections (e.g. QBE / Marsh), transcript excerpt, or any combination.\n"
        "There is no guaranteed structure or content order.\n\n"
        "IMPORTANT: client name, meeting date, and attendees will be sourced from the\n"
        "PowerPoint deck separately. Set client, meeting_date, and attendees to null / []\n"
        "in your output — do NOT attempt to infer them from this document.\n\n"
        "Extract the following:\n\n"
        "**1. METADATA** (best-effort — will be overridden by PPT):\n"
        "subject/meeting title (if clearly stated), location, author/note-taker.\n"
        "Set client=null, meeting_date=null, attendees=[].\n\n"
        "**2. DISCUSSION POINTS** — group related content under topic headings:\n"
        "- Explicit headings → use them as topic names exactly as written.\n"
        "- Party/speaker labels (e.g. QBE:, Marsh:, AIG:) → each party = one topic.\n"
        "- No headings → group thematically into 3–6 broad topics\n"
        "  (e.g. 'Carrier Update', 'Marsh Update', 'Market Conditions', 'Performance').\n"
        "- Prefer fewer broader topics. Copy bullet text VERBATIM. Preserve all numbers.\n"
        "- IMPORTANT: Remove any speaker attribution prefix from bullets.\n"
        "  e.g. 'Vinay Makwana outlined...' → 'Marsh outlined...'\n"
        "  e.g. 'John Smith said rates are up' → 'Rates are up'\n"
        "  Rewrite in plain factual third-person. Do not lose the substance.\n\n"
        "**3. ACTION ITEMS** (table, list, or inline):\n"
        "- action: the action description only — strip any speaker prefix AND remove any inline\n"
        "  timeline/date reference from the action text itself.\n"
        "  e.g. 'Arrange a roundtable for the energy sector in Q2 2026' -> action = 'Arrange a roundtable for the energy sector'\n"
        "  e.g. 'Share updated pricing by end of H1 2025' -> action = 'Share updated pricing'\n"
        "- owner: the person or team responsible (null if not stated).\n"
        "- due_date: extract any timeline or date reference — quarter (Q1/Q2/Q3/Q4 YYYY),\n"
        "  half-year (H1/H2 YYYY), month, year, or phrase like 'by end of year', 'next QBR'.\n"
        "  If stated in the action text, move it here and remove it from the action field.\n"
        "  Set to null only if no timeline is mentioned anywhere.\n\n"
        "RULES: Do not invent or rephrase substance. Missing fields → null or [].\n\n"
        "Return ONLY valid JSON with exactly these keys:\n"
        '{"client": null, "subject": null, "meeting_date": null, "location": null, '
        '"author": null, "attendees": [], '
        '"discussion_points": [{"topic": "...", "bullets": ["..."]}], '
        '"action_items": [{"action": "...", "owner": null, "due_date": null}]}\n\n'
        "Meeting note content:\n" + body
    )


def parse_meeting_note(pages: List[NotePage], call_json: JsonCaller) -> dict:
    """One LLM call: pages in, meeting-note structure out."""
    return call_json(build_parse_prompt(pages), label="parse_meeting_note", phase="notes")


def run(note_path: Path, output_path: Path, call_json: JsonCaller) -> dict:
    """Read the note, parse it, save the JSON, return the meeting data."""
    pages = read_note(note_path)
    log.info("MoM: read %d page(s) from %s", len(pages), Path(note_path).name)

    meeting_data = parse_meeting_note(pages, call_json)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(meeting_data, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return meeting_data
