"""The minutes as a Word document.

Title block, agenda, body, action items, page numbers — in that order, every time.
Only the body differs between the two modes, so that one step is a dispatch table and
everything around it is written once.

Pure output: it takes the meeting data and the written summary and produces a file. It
makes no decisions about content and calls no model.
"""
from __future__ import annotations

from pathlib import Path
from typing import Callable, Dict, List, Sequence

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from logger import get_logger

log = get_logger(__name__)

# Marsh navy — the same value as `--va-navy` in assets/theme_tokens.css.
NAVY = RGBColor(0x00, 0x0F, 0x47)
NAVY_HEX = "#000F47"
GREY = RGBColor(0x7B, 0x79, 0x74)

NOTHING_FOUND = "No material content identified."


# ── low-level helpers ─────────────────────────────────────────────────────────


def xml_safe(text: str) -> str:
    """Strip XML-illegal control characters and collapse whitespace."""
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    cleaned = re.sub(r"[\r\n]+", " ", cleaned)
    return re.sub(r" {2,}", " ", cleaned).strip()


def _centred(document: Document, text: str, *, size: int, bold=False, italic=False,
             color: RGBColor | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = NAVY


def _bullets(document: Document, items: Sequence[str]) -> None:
    for item in items:
        text = xml_safe(item)
        if text:
            document.add_paragraph(text, style="List Bullet")


def _shade(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_numbers(document: Document) -> None:
    """A centred PAGE field in the footer."""
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for tag, value in (("fldChar", "begin"), ("instrText", "PAGE"), ("fldChar", "end")):
        element = OxmlElement(f"w:{tag}")
        if tag == "instrText":
            element.text = f" {value} "
        else:
            element.set(qn("w:fldCharType"), value)
        run._r.append(element)


# ── the shared sections ───────────────────────────────────────────────────────


def _write_title_block(document: Document, meeting_data: dict, top_pairs: Sequence[dict]) -> str:
    """Title, date, client, attendees, priority topics. Returns the client name."""
    client = meeting_data.get("client") or "QBR Summary"
    subject = meeting_data.get("subject") or client
    date = meeting_data.get("meeting_date") or ""
    attendees = meeting_data.get("attendees") or []

    _centred(document, subject, size=18, bold=True, color=NAVY)
    _centred(document, f"Meeting Summary{' — ' + date if date else ''}", size=11, italic=True)

    if client and client.lower() not in subject.lower():
        _centred(document, f"Client: {client}", size=10)

    if attendees:
        shown = "  |  ".join(attendees[:10])
        if len(attendees) > 10:
            shown += f"  + {len(attendees) - 10} more"
        _centred(document, f"Attendees: {shown}", size=9)

    if top_pairs:
        labels = [f"{p['umbrella_tag']} / {p['sub_tag']}" for p in top_pairs[:5]]
        _centred(document, "Priority topics: " + "  ·  ".join(labels), size=8, color=GREY)

    document.add_paragraph()
    return client


def agenda_items(meeting_data: dict, ppt_sections: Sequence[dict]) -> List[str]:
    """The agenda, from the deck's own agenda slide, else its sections, else the topics."""
    from_deck = meeting_data.get("ppt_agenda_items") or []
    if from_deck:
        return list(from_deck)

    if ppt_sections:
        return [
            section["section_title"]
            for section in ppt_sections
            if section.get("section_title", "").lower() not in ("introduction", "intro")
            and section.get("section_number", 1) != 1
        ]

    return [
        point["topic"]
        for point in meeting_data.get("discussion_points", [])
        if point.get("topic")
    ]


def _write_agenda(document: Document, meeting_data: dict, ppt_sections: Sequence[dict]) -> None:
    _heading(document, "Agenda")
    _bullets(document, agenda_items(meeting_data, ppt_sections))


def _write_action_items(document: Document, meeting_data: dict) -> None:
    actions = [a for a in meeting_data.get("action_items", []) if a.get("action", "").strip()]
    if not actions:
        return

    _heading(document, "Action Items")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    for index, name in enumerate(["Action", "Action By", "Timeline"]):
        cell = table.rows[0].cells[index]
        cell.text = name
        _shade(cell, NAVY_HEX)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for action in actions:
        cells = table.add_row().cells
        for index, key in enumerate(("action", "owner", "due_date")):
            cells[index].text = action.get(key) or ""
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


# ── the body, one writer per mode ─────────────────────────────────────────────


def _write_section(document: Document, heading: str, bullets: Sequence[str]) -> None:
    _heading(document, heading)
    if bullets and list(bullets) != [NOTHING_FOUND]:
        _bullets(document, bullets)
    else:
        document.add_paragraph(NOTHING_FOUND).italic = True


def _write_skill_body(document: Document, summary: dict, client: str) -> None:
    for heading, key in (
        ("Strategy & Initiatives", "strategy_and_initiatives"),
        ("Country / Product / Region", "country_product_region"),
        ("Key Takeaways", "key_takeaways"),
    ):
        _write_section(document, heading, summary.get(key, []))


def _write_carrier_marsh_body(document: Document, summary: dict, client: str) -> None:
    carrier_heading = client.strip() if client and client.strip() else "Carrier Update"
    for heading, key in (
        (carrier_heading, "carrier_update"),
        ("Marsh Update", "marsh_update"),
    ):
        _write_section(document, heading, summary.get(key, []))


BODY_WRITERS: Dict[str, Callable[[Document, dict, str], None]] = {
    "skill": _write_skill_body,
    "carrier_marsh": _write_carrier_marsh_body,
}


# ── entry point ───────────────────────────────────────────────────────────────


def write_minutes(
    *,
    meeting_data: dict,
    summary: dict,
    top_pairs: Sequence[dict],
    ppt_sections: Sequence[dict],
    mode: str,
    path: Path,
) -> Path:
    """Write the document and return the path it was saved to."""
    write_body = BODY_WRITERS.get(mode)
    if write_body is None:
        raise ValueError(f"mode must be one of {tuple(BODY_WRITERS)}, got {mode!r}")

    document = Document()
    for section in document.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1.25)

    client = _write_title_block(document, meeting_data, top_pairs)
    _write_agenda(document, meeting_data, ppt_sections)
    write_body(document, summary, client)
    _write_action_items(document, meeting_data)
    _add_page_numbers(document)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    log.info("MoM: wrote %s", path.name)
    return path
