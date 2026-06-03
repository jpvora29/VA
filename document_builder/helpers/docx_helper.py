import os
import re
from datetime import datetime
from typing import Any, Optional
from collections import defaultdict
from dataclasses import dataclass, fields

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import io

from document_builder.helpers.xml_helper import *
from config.report_config import (
    NAVY,
    ELECTRIC_BLUE,
    LIGHT_BLUE,
    WHITE,
    DARK,
    GREEN_TEXT,
    RED_TEXT,
    LIGHT_GRAY,
    GRAY,
)
from document_builder.models import TopKPIs
from document_builder.helpers.number_formatter import format_money, format_pct
from document_builder.helpers.table_pivot_helper import *
from document_builder.helpers.field_identification import find_field, classify_question


def _rows_from_query_result(query_result: Any) -> list[dict[str, Any]]:
    if not query_result:
        return []
    if isinstance(query_result, list):
        return [row for row in query_result if isinstance(row, dict)]
    if isinstance(query_result, dict):
        rows: list[dict[str, Any]] = []
        for lens, value in query_result.items():
            if isinstance(value, list):
                for row in value:
                    enriched = dict(row)
                    enriched.setdefault("Source", str(lens).title())
                    rows.append(enriched)
            elif isinstance(value, dict):
                enriched = dict(value)
                enriched.setdefault("Source", str(lens).title())
                rows.append(enriched)
        return rows
    return []


def run(
    paragraph: Paragraph,
    text: str,
    bold: bool = False,
    size_pt: float = 10,
    color: RGBColor = None,
    name: str = "Arial",
) -> Any:
    r = paragraph.add_run(text)
    r.font.name = name
    r.font.size = Pt(size_pt)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def add_labeled_paragraph(document: DocumentType, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"{label}:")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 44, 119)
    paragraph.add_run(value or "Not provided")


def add_section_title(
    doc: DocumentType, title: str, accent: Optional[RGBColor] = None
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_paragraph_bottom_rule(p, color="D0D8E8", size=4)
    run(p, title, bold=True, size_pt=13, color=accent or NAVY)


def add_section_caption(doc: DocumentType, text: str) -> None:
    """One-line gray caption under a section title."""
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    run(cap, text, size_pt=9, color=GRAY)


def style_simple_table(table: Table, header_fill: str = "#000F47") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
            set_cell_borders(cell, {"bottom": "D8DFF0"})
            if row_index == 0:
                goa(cell._tc, "w:tcPr").append(shading(header_fill))
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = WHITE
                        run.bold = True
                        run.font.size = Pt(8.5)
            else:
                goa(cell._tc, "w:tcPr").append(
                    shading("#F7F9FD" if row_index % 2 == 0 else "#FFFFFF")
                )
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = DARK
                        run.font.size = Pt(8.8)


def add_kpi_title(
    cell: _Cell,
    label: str,
    value: str,
    delta_text: str = "",
    delta_good: Optional[bool] = None,
) -> None:
    goa(cell._tc, "w:tcPr").append(shading("EEF2F8"))
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    set_cell_borders(cell, {"bottom": "0A4AB0"})
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(2)
    run(p_label, label.upper(), bold=True, size_pt=7, color=GRAY)

    p_value = cell.add_paragraph()
    p_value.paragraph_format.space_before = Pt(0)
    p_value.paragraph_format.space_after = Pt(2)
    run(p_value, value or "-", bold=True, size_pt=16, color=NAVY)

    if delta_text:
        p_delta = cell.add_paragraph()
        p_delta.paragraph_format.space_before = Pt(0)
        p_delta.paragraph_format.space_after = Pt(0)
        if delta_good is True:
            color = GREEN_TEXT
        elif delta_good is False:
            color = RED_TEXT
        else:
            color = GRAY

        run(p_delta, delta_text, bold=True, size_pt=8, color=color)


def add_kpi_strip(doc: DocumentType, kpis: TopKPIs) -> None:

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_table_full_width(table)

    # Total Premium
    add_kpi_title(
        table.cell(0, 0),
        "Total Premium",
        format_money(kpis.total_premium) if kpis.total_premium else "-",
    )

    # YoY %
    yoy = kpis.yoy
    add_kpi_title(
        table.cell(0, 1),
        "YoY Growth",
        format_pct(yoy) if yoy is not None else "-",
        delta_text=(
            "▲ vs prior year"
            if (yoy or 0) > 0
            else ("▼ vs prior year" if (yoy or 0) < 0 else "")
        ),
        delta_good=(yoy is not None and yoy > 0) if yoy is not None else None,
    )

    # Rank
    rank = kpis.gpr_rank
    rank_delta = kpis.rank_delta
    rank_delta_text = ""
    rank_delta_good = None
    if rank_delta is not None and rank_delta != 0:
        arrow = "▲" if rank_delta < 0 else "▼"
        rank_delta_text = f"{arrow} {abs(rank_delta):.0f} vs prior"
        rank_delta_good = rank_delta < 0

    add_kpi_title(
        table.cell(0, 2),
        "GPR Rank",
        f"#{int(rank)}" if rank is not None else "-",
        delta_text=rank_delta_text,
        delta_good=rank_delta_good,
    )

    # Score
    score = kpis.survey_score
    if score is not None:
        add_kpi_title(table.cell(0, 3), "Survey Score", f"{score}")
    else:
        add_kpi_title(table.cell(0, 3), "Survey Score", "-")


def add_question_subtitle(doc: DocumentType, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.05)
    run(p, text, bold=True, size_pt=10, color=ELECTRIC_BLUE)


def add_pivot_table(
    doc: DocumentType,
    headers: list[str],
    rows: list[list[Any]],
    delta_cols: list[tuple[int, bool]],
) -> None:
    if not headers or not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell.text = str(row_values[idx]) if idx < len(row_values) else ""
    style_simple_table(table, header_fill="#000F47")

    for col_idx, good_when_pos in delta_cols:
        for r_idx in range(1, len(table.rows)):
            apply_delta_shading(
                table.cell(r_idx, col_idx), None, good_when_positive=good_when_pos
            )


def add_product_breakdown_section(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    filters: dict[str, Any],
) -> None:
    if not question_results:
        return

    candidates: list[tuple[dict[str, Any], list[dict[str, Any]], str]] = []
    for r in question_results:
        rows = _rows_from_query_result(r.get("pitch_query_result"))
        if not rows:
            continue
        if not any(find_field(row, "product") for row in rows):
            continue
        candidates.append((r, rows, classify_question(r.get("question", ""))))

    if not candidates:
        return

    add_section_title(doc, "Product Breakdown")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    run(
        intro,
        "Product-level detail behind each performance question. Green / red shading indicates "
        "directionally favourable or unfavourable movement.",
        size_pt=9,
        color=GRAY,
    )

    try:
        target_year = int(filters.get("year", None))
    except (TypeError, ValueError):
        target_year = None

    for result, rows, qtype in candidates:
        if qtype in ("appetite", "peer", "rank"):
            headers, pivot_rows, delta_cols = rows_to_styled_table(
                result.get("pitch_query_result") or ""
            )
            if not pivot_rows:
                continue
            add_question_subtitle(doc, result.get("question") or "")
            add_pivot_table(doc, headers, pivot_rows, delta_cols)


# --- Whitespace / Industry / Segment analysis sections -----------------------


def _question_matches(question: str, keywords: tuple[str, ...]) -> bool:
    q = (question or "").lower()
    return any(k in q for k in keywords)


def _render_question_tables(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    *,
    row_field: Optional[str],
    keywords: tuple[str, ...],
) -> int:
    """Render a styled pivot table for each question result that belongs to this
    section — matched either by a dimension present in its rows (`row_field`) or by
    its question text (`keywords`). Returns how many tables were rendered.
    """
    rendered = 0
    for result in question_results:
        rows = _rows_from_query_result(result.get("pitch_query_result"))
        if not rows:
            continue
        question = result.get("question") or ""
        has_field = bool(row_field) and any(find_field(row, row_field) for row in rows)
        if not (has_field or _question_matches(question, keywords)):
            continue

        headers, pivot_rows, delta_cols = rows_to_styled_table(
            result.get("pitch_query_result") or ""
        )
        if not pivot_rows:
            continue
        add_question_subtitle(doc, question)
        add_pivot_table(doc, headers, pivot_rows, delta_cols)
        rendered += 1
    return rendered


def _section_has_content(
    question_results: list[dict[str, Any]],
    *,
    row_field: Optional[str],
    keywords: tuple[str, ...],
) -> bool:
    """True if any question result has rows belonging to this section."""
    for r in question_results:
        rows = _rows_from_query_result(r.get("pitch_query_result"))
        if not rows:
            continue
        if row_field and any(find_field(row, row_field) for row in rows):
            return True
        if _question_matches(r.get("question") or "", keywords):
            return True
    return False


def _add_analysis_section(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    *,
    title: str,
    caption: str,
    row_field: Optional[str],
    keywords: tuple[str, ...],
    accent: Optional[RGBColor] = None,
) -> None:
    """Emit a titled analysis section only when there is grounded content."""
    if not question_results or not _section_has_content(
        question_results, row_field=row_field, keywords=keywords
    ):
        return
    add_section_title(doc, title, accent=accent)
    add_section_caption(doc, caption)
    _render_question_tables(
        doc, question_results, row_field=row_field, keywords=keywords
    )


def add_whitespace_section(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    accent: Optional[RGBColor] = None,
) -> None:
    """Slices where the carrier is absent/thin while the Marsh book participates."""
    _add_analysis_section(
        doc,
        question_results,
        title="Whitespace Analysis",
        caption=(
            "Slices where the carrier has zero or materially thin premium while the "
            "Marsh book participates — portfolio gaps and growth headroom."
        ),
        row_field=None,
        keywords=("whitespace", "white space"),
        accent=accent,
    )


def add_industry_section(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    accent: Optional[RGBColor] = None,
) -> None:
    """Performance by SIC major/minor industry class."""
    _add_analysis_section(
        doc,
        question_results,
        title="Industry-Level Analysis",
        caption=(
            "Premium, growth, and share by industry (SIC class). Green / red shading "
            "marks directionally favourable or unfavourable movement."
        ),
        row_field="industry",
        keywords=("industry", "sic", "sector"),
        accent=accent,
    )


def add_segment_section(
    doc: DocumentType,
    question_results: list[dict[str, Any]],
    accent: Optional[RGBColor] = None,
) -> None:
    """Performance by client/business segment."""
    _add_analysis_section(
        doc,
        question_results,
        title="Segment Analysis",
        caption=(
            "Premium, growth, and share by client segment. Green / red shading marks "
            "directionally favourable or unfavourable movement."
        ),
        row_field="segment",
        keywords=("segment",),
        accent=accent,
    )
