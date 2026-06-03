from docx.document import Document as DocumentType
from docx.shared import Inches, Pt

from document_builder.helpers.docx_helper import add_paragraph_bottom_rule, run
from document_builder.helpers.design_spec import load_design_spec
from config.report_config import NAVY, GRAY


def add_section_heading(doc: DocumentType) -> None:
    spec = load_design_spec()
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(2)
    sp2.paragraph_format.space_after = Pt(8)

    sec_heading = doc.add_paragraph()
    sec_heading.paragraph_format.space_before = Pt(0)
    sec_heading.paragraph_format.space_after = Pt(2)
    add_paragraph_bottom_rule(sec_heading, color=spec.hex("navy"), size=8)
    run(sec_heading, "▌ ", bold=True, size_pt=13, color=NAVY)
    run(sec_heading, "Executive Narrative", bold=True, size_pt=13, color=NAVY)

    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    run(
        caption,
        "Senior-consultant read of performance, positioning, and the strategic so-what.",
        size_pt=9,
        color=GRAY,
    )
