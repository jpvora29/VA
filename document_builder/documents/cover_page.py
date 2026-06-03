"""Cover page — the report's first page.

A clean, centered typographic cover: thin Marsh navy/red accent bands top and
bottom, a large title in the heading font, the carrier in electric blue, the
country/year meta, a hairline divider, the prepared date, and a confidentiality
line. Ends with a page break so the body starts on page 2 (which carries the
running header/footer). All colours/labels/sizes come from the design skill via
`load_design_spec()`.
"""
from __future__ import annotations

from datetime import datetime

from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config.report_config import HEADING_FONT, BODY_FONT
from document_builder.helpers.design_spec import load_design_spec
from document_builder.helpers.docx_helper import (
    run,
    set_cell_margins,
    set_row_height,
    set_table_full_width,
)
from document_builder.helpers.xml_helper import (
    add_paragraph_bottom_rule,
    remove_table_borders,
    shading,
)


def _accent_band(doc: DocumentType, fill_hex: str, height_twips: int) -> None:
    """A thin full-width colour band (single-cell borderless table)."""
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(band)
    set_table_full_width(band)
    cell = band.cell(0, 0)
    cell._tc.get_or_add_tcPr().append(shading(fill_hex))
    set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
    set_row_height(band.rows[0], height_twips)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _spacer(doc: DocumentType, pts: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


def _centered(doc: DocumentType, space_after: float = 0) -> "Any":  # noqa: F821
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_cover_page(
    doc: DocumentType, carrier: str, country: str, year: object
) -> None:
    spec = load_design_spec()
    navy = spec.rgb("navy")
    electric = spec.rgb("electric_blue")
    red = spec.rgb("red_text")
    gray = spec.rgb("gray")

    # Top accent bands.
    _accent_band(doc, spec.hex_hash("navy"), 60)
    _accent_band(doc, spec.hex_hash("red_text"), 24)

    # Push the title block toward the visual centre of the page.
    _spacer(doc, 210)

    p_kicker = _centered(doc, space_after=14)
    r = run(
        p_kicker,
        spec.label("kicker"),
        bold=True,
        size_pt=spec.sizes.get("cover_kicker", 11),
        color=gray,
        name=BODY_FONT,
    )
    r.font.all_caps = True

    p_title = _centered(doc, space_after=6)
    run(
        p_title,
        spec.label("title"),
        bold=True,
        size_pt=spec.sizes.get("cover_title", 40),
        color=navy,
        name=HEADING_FONT,
    )

    p_carrier = _centered(doc, space_after=4)
    run(
        p_carrier,
        str(carrier or "").upper(),
        bold=True,
        size_pt=spec.sizes.get("cover_subtitle", 20),
        color=electric,
        name=HEADING_FONT,
    )

    p_meta = _centered(doc, space_after=14)
    run(
        p_meta,
        f"{country}   •   {year}",
        size_pt=spec.sizes.get("cover_meta", 10) + 1,
        color=navy,
        name=BODY_FONT,
    )

    # Hairline divider centred under the title block.
    p_rule = _centered(doc, space_after=12)
    add_paragraph_bottom_rule(p_rule, color=spec.hex("rule_soft"), size=6)
    run(p_rule, " ", size_pt=2, color=gray)

    prepared = datetime.now().strftime("%d %B %Y")
    p_date = _centered(doc, space_after=2)
    run(p_date, f"Prepared {prepared}", size_pt=spec.sizes.get("cover_meta", 10), color=gray)

    p_conf = _centered(doc, space_after=0)
    run(
        p_conf,
        spec.label("confidentiality"),
        size_pt=spec.sizes.get("cover_meta", 10) - 1,
        color=gray,
        name=BODY_FONT,
    )

    # Bottom accent bands sit at the foot of the cover.
    _spacer(doc, 210)
    _accent_band(doc, spec.hex_hash("red_text"), 24)
    _accent_band(doc, spec.hex_hash("navy"), 60)

    doc.add_page_break()
