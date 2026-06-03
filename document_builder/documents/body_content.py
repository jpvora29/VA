import re

from docx.document import Document as DocumentType

from docx.shared import Inches, Pt

from document_builder.helpers.docx_helper import run, add_paragraph_bottom_rule
from document_builder.helpers.design_spec import load_design_spec

from config.report_config import SECTION_HEADINGS, ELECTRIC_BLUE, NAVY, DARK, GRAY


# "A. Premium Performance Perspective" / "B) Broker Perspective" — lettered sub-heads.
_LETTER_SUBHEAD = re.compile(r"^[A-Z][.)]\s+\S")


def _normalize_heading_text(line: str) -> str:
    """Strip markdown markers and normalize for heading detection."""
    s = re.sub(r"^#{1, 6}\s*", "", line)  # ## markers
    s = re.sub(r"\*{1, 3}(.*?)\*{1, 3}", r"\1", s)  # **bold** / *italic*
    s = re.sub(r"_{1,2}(.*?)_{1,2}", r"\1", s)  # __bold__
    s = re.sub(r"^\d+\.\s*", "", s)  # "1. "
    s = s.lower().strip().rstrip(":").strip()
    s = s.split(" - ")[0].rstrip(":").strip()
    return s


def _strip_markdown(line: str) -> str:
    """Return display text with markdown formatting removed."""
    s = re.sub(r"^#{1,6}\s*", "", line)
    s = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", s)
    s = re.sub(r"_{1,2}(.*?)_{1,2}", r"\1", s)
    return s.strip().rstrip(":")


def add_body_content(doc: DocumentType, report_summary: str | None) -> None:
    if not report_summary:
        return

    spec = load_design_spec()
    rule_soft = spec.hex("rule_soft")

    for line in report_summary.splitlines():
        raw = line.strip()
        if not raw:
            continue

        normalized = _normalize_heading_text(raw)
        is_heading = bool(re.match(r"^#{1,6}\s", raw)) or normalized in SECTION_HEADINGS
        is_subhead = (not is_heading) and bool(_LETTER_SUBHEAD.match(raw))
        is_bullet = raw.startswith(("-", "•", "*")) and not is_heading

        p = doc.add_paragraph()

        if is_heading:
            display = _strip_markdown(raw).title()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            add_paragraph_bottom_rule(p, color=rule_soft, size=6)
            run(p, "▌ ", bold=True, size_pt=12, color=NAVY)
            run(p, display, bold=True, size_pt=12, color=NAVY)

        elif is_subhead:
            display = _strip_markdown(raw)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.05)
            run(p, display, bold=True, size_pt=10.5, color=ELECTRIC_BLUE)

        elif is_bullet:
            bullet_text = re.sub(r"^[-•*]\s*", "", raw)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.line_spacing = 1.12
            run(p, "•  ", bold=True, size_pt=10, color=ELECTRIC_BLUE)
            run(p, bullet_text, size_pt=10, color=DARK)

        else:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.18
            run(p, raw, size_pt=10, color=DARK)
