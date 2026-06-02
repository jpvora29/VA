from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from document_builder.helpers.docx_helper import (
    remove_table_borders,
    set_table_full_width,
    set_cell_margins,
    set_cell_borders,
    set_row_height,
    run,
)

from document_builder.helpers.xml_helper import shading
from config.report_config import WHITE, HEADING_FONT


def _add_header_banner(doc: DocumentType) -> None:
    accent = doc.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(accent)
    set_table_full_width(accent)
    ac = accent.cell(0, 0)
    # ac._tc.get_or_add_tcPr().append(_shading("C53532"))
    set_cell_margins(ac, top=0, bottom=0, left=0, right=0)
    set_row_height(accent.rows[0], 72)
    p_acc = ac.paragraphs[0]
    p_acc.paragraph_format.space_before = Pt(0)
    p_acc.paragraph_format.space_after = Pt(0)


def _add_blue_banner(
    doc: DocumentType, *, carrier: str, country: str, year: str
) -> None:

    # Blue main banner
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(banner)
    set_table_full_width(banner)
    bc = banner.cell(0, 0)
    bc._tc.get_or_add_tcPr().append(shading("#000F47"))
    set_cell_margins(bc, top=280, bottom=260, left=280, right=280)

    p_kicker = bc.paragraphs[0]
    p_kicker.paragraph_format.space_before = Pt(0)
    p_kicker.paragraph_format.space_after = Pt(5)
    run(
        p_kicker,
        "MARSH ICG  - CARRIER REPORT",
        bold=False,
        size_pt=7,
        color=RGBColor(115, 150, 205),
        name=HEADING_FONT,
    )

    p_title = bc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(5)
    run(
        p_title,
        "Performance Analysis Builder",
        bold=True,
        size_pt=22,
        color=WHITE,
        name=HEADING_FONT,
    )

    p_sub = bc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    run(
        p_sub,
        f"• {carrier} • {country} • {year}",
        size_pt=9.5,
        color=RGBColor(155, 185, 230),
    )

    p_source = bc.add_paragraph()
    p_source.paragraph_format.space_before = Pt(0)
    p_source.paragraph_format.space_after = Pt(0)
    run(
        p_source,
        f"Source - GPR, Carrier Survey",
        size_pt=9.5,
        color=RGBColor(155, 185, 230),
    )


def add_header(doc: DocumentType, carrier: str, country: str, year: int) -> None:
    _add_header_banner(doc)
    _add_blue_banner(doc, carrier=carrier, country=country, year=year)
