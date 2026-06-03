"""Running page header + footer (page 2 onward).

Header: `banner_title` left in navy, `carrier • country • year` right in gray,
over a soft hairline rule. Footer: confidentiality left, `Page X of Y` centred
(Word field codes), carrier right. The cover page (page 1) is excluded via
`different_first_page_header_footer`. Styling/labels come from the design skill.
"""
from __future__ import annotations

from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from config.report_config import BODY_FONT, CONTENT_WIDTH
from document_builder.helpers.design_spec import load_design_spec
from document_builder.helpers.docx_helper import run, set_cell_margins
from document_builder.helpers.xml_helper import (
    add_paragraph_bottom_rule,
    remove_table_borders,
    set_table_full_width,
)


def _add_field(paragraph, instr: str, *, size_pt: float, color, bold: bool = False):
    """Append a Word field (e.g. PAGE / NUMPAGES) styled like a normal run."""
    r = paragraph.add_run()
    r.font.name = BODY_FONT
    r.font.size = Pt(size_pt)
    r.font.color.rgb = color
    r.bold = bold
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr_el)
    r._r.append(end)


def _build_header(section, carrier: str, country: str, year: object) -> None:
    spec = load_design_spec()
    header = section.header
    header.is_linked_to_previous = False
    # Reuse the auto-created empty paragraph for the rule; build content in a table.
    base = header.paragraphs[0]
    base.paragraph_format.space_after = Pt(2)

    table = header.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_table_full_width(table)
    left, right = table.rows[0].cells
    for c in (left, right):
        set_cell_margins(c, top=10, bottom=40, left=0, right=0)

    pl = left.paragraphs[0]
    pl.paragraph_format.space_after = Pt(0)
    r = run(pl, spec.label("banner_title"), bold=True, size_pt=spec.sizes.get("page_header", 8), color=spec.rgb("navy"))

    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_after = Pt(0)
    meta = "  •  ".join(str(x) for x in (carrier, country, year) if x)
    run(pr, meta, size_pt=spec.sizes.get("page_header", 8), color=spec.rgb("gray"))

    # Soft rule under the header row.
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(1)
    rule_p.paragraph_format.space_after = Pt(0)
    add_paragraph_bottom_rule(rule_p, color=spec.hex("rule_soft"), size=6)


def _build_footer(section, carrier: str) -> None:
    spec = load_design_spec()
    footer = section.footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=3, width=CONTENT_WIDTH)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_table_full_width(table)
    left, center, right = table.rows[0].cells
    for c in (left, center, right):
        set_cell_margins(c, top=40, bottom=10, left=0, right=0)

    size = spec.sizes.get("page_footer", 8)
    gray = spec.rgb("gray")

    pl = left.paragraphs[0]
    pl.paragraph_format.space_after = Pt(0)
    run(pl, spec.label("confidentiality"), size_pt=size, color=gray)

    pc = center.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(0)
    run(pc, "Page ", size_pt=size, color=gray)
    _add_field(pc, "PAGE", size_pt=size, color=gray, bold=True)
    run(pc, " of ", size_pt=size, color=gray)
    _add_field(pc, "NUMPAGES", size_pt=size, color=gray, bold=True)

    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_after = Pt(0)
    run(pr, str(carrier or ""), size_pt=size, color=gray)


def add_page_furniture(
    doc: DocumentType, carrier: str, country: str, year: object
) -> None:
    """Attach the running header + footer; keep them off the cover (page 1)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    _build_header(section, carrier, country, year)
    _build_footer(section, carrier)
